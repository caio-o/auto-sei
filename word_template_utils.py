import os
import re
from datetime import datetime
from docx import Document


PLACEHOLDER_REGEX = re.compile(r"\{([a-zA-Z0-9_]+)\}")


def extract_placeholders_from_docx(docx_path: str) -> list[str]:
    doc = Document(docx_path)
    found = set()

    def collect(text: str):
        for match in PLACEHOLDER_REGEX.findall(text or ""):
            found.add(match.strip())

    for p in doc.paragraphs:
        collect(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                collect(cell.text)

    return sorted(found)


def normalize_context(data: dict) -> dict:
    ctx = dict(data)

    if not ctx.get("data_atual"):
        ctx["data_atual"] = datetime.now().strftime("%d/%m/%Y")

    for k, v in list(ctx.items()):
        if v is None:
            ctx[k] = ""

    return ctx


def _replace_text_in_paragraph(paragraph, context: dict):
    if not paragraph.text:
        return

    original = paragraph.text
    new_text = original

    for key, value in context.items():
        placeholder = "{" + str(key) + "}"
        new_text = new_text.replace(placeholder, str(value))

    if new_text == original:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


# 🔥 NOVO: aplica negrito em frases
def apply_bold_phrases(doc, frases: list[str]):
    for p in doc.paragraphs:
        texto = p.text

        for frase in frases:
            if not frase:
                continue

            if frase not in texto:
                continue

            # limpa o parágrafo
            for run in p.runs:
                run.text = ""

            partes = texto.split(frase)

            for i, parte in enumerate(partes):
                if parte:
                    p.add_run(parte)

                if i < len(partes) - 1:
                    r = p.add_run(frase)
                    r.bold = True


def render_template_to_docx(template_path: str, output_path: str, context: dict):
    context = normalize_context(context)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Modelo não encontrado: {template_path}")

    pasta_destino = os.path.dirname(output_path)
    if pasta_destino:
        os.makedirs(pasta_destino, exist_ok=True)

    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except PermissionError:
            raise PermissionError(
                f"Não foi possível substituir o arquivo porque ele está aberto:\n{output_path}"
            )

    doc = Document(template_path)

    # 🔹 substitui placeholders
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, context)

    # AQUI ENTRA O NEGRITO

    # versão dinâmica (RECOMENDADA)
    frase_dinamica = (
        f"emitiu-se o extrato de prestações ({context.get('extrato')}), "
        f"a memória de cálculo ({context.get('memoria')}) "
        f"e a GRU ({context.get('gru')})"
    )

    frases_negrito = [frase_dinamica]

    apply_bold_phrases(doc, frases_negrito)

    # salva
    doc.save(output_path)